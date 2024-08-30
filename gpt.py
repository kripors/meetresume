import google.generativeai as genai
import os
import json
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist

# Baixar recursos necessários do NLTK
nltk.download('punkt')
nltk.download('stopwords')

# Carregar a configuração do arquivo JSON
def carregar_configuracoes(caminho_arquivo):
    with open(caminho_arquivo, 'r') as arquivo:
        return json.load(arquivo)

# Carregar a configuração
config = carregar_configuracoes('config.json')
api_key = config.get('GOOGLE_API_KEY')
caminho_base = config.get('caminho_base')

if api_key is None:
    raise ValueError("A chave da API 'GOOGLE_API_KEY' não está definida no arquivo de configuração.")

# Configurar a API do Google (Gemini)
genai.configure(api_key=api_key)

# Configurar o modelo
model = genai.GenerativeModel('gemini-1.5-flash')

# Definindo os caminhos de entrada e saída
caminho_entrada = f"{caminho_base}.txt"
caminho_saida = f"{caminho_base}.docx"

def ler_arquivo_txt(caminho_arquivo):
    with open(caminho_arquivo, 'r', encoding='utf-8') as arquivo:
        return arquivo.read()

def preprocessar_texto(texto):
    # Tokenização de sentenças e palavras
    sentencas = sent_tokenize(texto)
    palavras = word_tokenize(texto.lower())
    
    # Remover stopwords
    stop_words = set(stopwords.words('portuguese'))
    palavras_filtradas = [palavra for palavra in palavras if palavra.isalnum() and palavra not in stop_words]
    
    # Análise de frequência de palavras
    freq_dist = FreqDist(palavras_filtradas)
    palavras_mais_comuns = freq_dist.most_common(10)
    
    return sentencas, palavras_filtradas, palavras_mais_comuns


def gerar_resumo_profissional(conteudo):
    sentencas, palavras_filtradas, palavras_mais_comuns = preprocessar_texto(conteudo)

    prompt = f"""
    Como professor de português, analise o conteúdo do arquivo em anexo e faça um texto destacando os principais pontos, com um tom profissional.

    Conteúdo original:
    {conteudo}

    Análise prévia:
    - Número de sentenças: {len(sentencas)}
    - Palavras-chave mais frequentes: {', '.join([palavra for palavra, _ in palavras_mais_comuns])}

    O resumo deve:
    1. Ter um tom profissional e formal
    2. Destacar os pontos-chave do texto original
    3. Ser estruturado em tópicos claros
    4. Usar a seguinte formatação:
       - Tópicos principais sem marcadores
       - Subtópicos com • (ex: • Subtópico)
       - Sub-subtópicos com dois espaços e • (ex:   • Sub-subtópico)
    5. Manter uma estrutura hierárquica clara
    6. NÃO usar asteriscos (**) ou hashtags (###) na formatação

    Por favor, formate o resumo de maneira clara e legível, seguindo estritamente as regras de formatação acima.
    """

    resposta = model.generate_content(prompt)
    return resposta.text

def salvar_como_doc(conteudo, caminho_arquivo):
    documento = Document()

    # Definir estilos
    estilos = {
        'Normal': {
            'nome': 'Calibri',
            'tamanho': 11,
            'alinhamento': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        },
        'Titulo': {
            'nome': 'Calibri',
            'tamanho': 12,
            'negrito': True,
            'alinhamento': WD_PARAGRAPH_ALIGNMENT.LEFT
        },
        'Marcador1': {
            'nome': 'Calibri',
            'tamanho': 11,
            'alinhamento': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'estilo': 'List Bullet'
        },
        'Marcador2': {
            'nome': 'Calibri',
            'tamanho': 11,
            'alinhamento': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'estilo': 'List Bullet 2'
        }
    }

    # Criar ou atualizar estilos
    for nome, config in estilos.items():
        if nome in documento.styles:
            estilo = documento.styles[nome]
        else:
            estilo = documento.styles.add_style(nome, WD_STYLE_TYPE.PARAGRAPH)
        
        estilo.font.name = config['nome']
        estilo.font.size = Pt(config['tamanho'])
        if 'negrito' in config:
            estilo.font.bold = config['negrito']
        if 'estilo' in config:
            estilo.base_style = documento.styles[config['estilo']]

    # Processar o conteúdo
    linhas = conteudo.split('\n')
    nivel_atual = 0

    for linha in linhas:
        linha = linha.strip()
        if not linha:
            continue

        # Remover asteriscos e hashtags
        linha = linha.replace('*', '').replace('#', '')

        if not linha.startswith('•'):
            # Tópico principal
            p = documento.add_paragraph(linha, style='Titulo')
            nivel_atual = 0
        elif linha.startswith('•'):
            # Marcador de primeiro nível
            p = documento.add_paragraph(linha.lstrip('• '), style='Marcador1')
            nivel_atual = 1
        elif linha.startswith('  •'):
            # Marcador de segundo nível
            p = documento.add_paragraph(linha.lstrip(' •'), style='Marcador2')
            nivel_atual = 2
        else:
            # Texto normal
            p = documento.add_paragraph(linha, style='Normal')

        p.alignment = estilos[p.style.name]['alinhamento']

    documento.save(caminho_arquivo)



# Execução principal
conteudo_original = ler_arquivo_txt(caminho_entrada)
resumo_profissional = gerar_resumo_profissional(conteudo_original)

salvar_como_doc(resumo_profissional, caminho_saida)

print(f"O resumo profissional foi gerado e salvo em {caminho_saida}")
print("\nConteúdo do resumo:")
print(resumo_profissional)